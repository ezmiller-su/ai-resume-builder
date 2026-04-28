from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import os


def set_font(run, name="Times New Roman", size=Pt(12), bold=False, color=RGBColor(0, 0, 0)):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    # Ensure Times New Roman works for East Asian / CS fonts too
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _get_or_create_pPr(paragraph):
    """Get or create pPr element for a paragraph."""
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        paragraph._element.insert(0, pPr)
    return pPr


# Schema-compliant ordering for w:pPr children
_PPR_ORDER = [
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
    "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
    "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens",
    "w:kinsoku", "w:wordWrap", "w:overflowPunct", "w:topLinePunct",
    "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing",
    "w:mirrorIndents", "w:suppressOverlap", "w:jc", "w:textDirection",
    "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl", "w:divId",
    "w:cnfStyle", "w:rPr",
]


def _insert_pPr_child(pPr, child_elem):
    """Insert a child element into pPr in schema-compliant order."""
    child_tag = child_elem.tag
    child_local = child_tag.split("}")[-1] if "}" in child_tag else child_tag
    child_qname = f"w:{child_local}"

    try:
        child_idx = _PPR_ORDER.index(child_qname)
    except ValueError:
        pPr.append(child_elem)
        return

    # Find the first existing sibling that should come after this element
    for existing in pPr:
        existing_tag = existing.tag
        existing_local = existing_tag.split("}")[-1] if "}" in existing_tag else existing_tag
        existing_qname = f"w:{existing_local}"
        try:
            existing_idx = _PPR_ORDER.index(existing_qname)
        except ValueError:
            continue
        if existing_idx > child_idx:
            existing.addprevious(child_elem)
            return

    pPr.append(child_elem)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_rule=None):
    """Set paragraph spacing."""
    pPr = _get_or_create_pPr(paragraph)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f"<w:spacing {nsdecls('w')}/>")
        _insert_pPr_child(pPr, spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)


def add_bottom_border(paragraph):
    """Add a bottom border to a paragraph (section header underline)."""
    pPr = _get_or_create_pPr(paragraph)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    _insert_pPr_child(pPr, pBdr)


def add_right_tab(paragraph, pos=10710):
    """Add a right-aligned tab stop."""
    pPr = _get_or_create_pPr(paragraph)
    tabs = parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        f'  <w:tab w:val="right" w:pos="{pos}"/>'
        f'</w:tabs>'
    )
    _insert_pPr_child(pPr, tabs)


def add_tab_run(paragraph):
    """Add a tab character run to a paragraph."""
    r = parse_xml(f'<w:r {nsdecls("w")}><w:tab/></w:r>')
    paragraph._element.append(r)
    return r


def make_bullet_paragraph(doc):
    """Create a bullet-pointed paragraph using Word's native bullet numbering."""
    p = doc.add_paragraph()
    pPr = _get_or_create_pPr(p)

    # Set as ListParagraph style
    pStyle = parse_xml(f'<w:pStyle {nsdecls("w")} w:val="ListParagraph"/>')
    _insert_pPr_child(pPr, pStyle)

    # Add numbering reference
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="0"/>'
        f'  <w:numId w:val="1"/>'
        f'</w:numPr>'
    )
    _insert_pPr_child(pPr, numPr)

    return p


def add_empty_paragraph(doc):
    """Add an empty paragraph with Times New Roman formatting."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    # Add an empty run to set the font
    r = p.add_run()
    set_font(r, bold=False)
    return p


def add_section_header(doc, text):
    """Add a section header with bottom border (e.g., EDUCATION, PROJECTS)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_bottom_border(p)
    r = p.add_run(text)
    set_font(r, bold=True)
    return p


def add_entry_header(doc, bold_text, normal_text, date_text):
    """Add an entry line with bold title, normal org, and right-tab date."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_right_tab(p, 10710)

    r1 = p.add_run(bold_text)
    set_font(r1, bold=True)

    r2 = p.add_run(normal_text)
    set_font(r2, bold=False)

    # Add tab + date in the same run
    tab_r = add_tab_run(p)
    # Set font on the tab run
    rPr = parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'  <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'  <w:color w:val="000000"/>'
        f'</w:rPr>'
    )
    tab_r.insert(0, rPr)

    r3 = p.add_run(date_text)
    set_font(r3, bold=False)

    return p


def add_bullet(doc, text):
    """Add a bullet point item."""
    p = make_bullet_paragraph(doc)
    set_paragraph_spacing(p, before=0, after=0)
    r = p.add_run(text)
    set_font(r, bold=False)
    return p


def build_resume():
    doc = Document()

    # -- Set default font in styles --
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Also set the element-level rFonts for the Normal style
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")

    # -- Set up ListParagraph style --
    list_style = doc.styles["List Paragraph"]
    list_style.font.name = "Times New Roman"
    list_style.font.size = Pt(12)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(0)

    # -- Page setup: US Letter, 0.5" margins --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    # Set section type to continuous
    sectPr = section._sectPr
    sect_type = sectPr.find(qn("w:type"))
    if sect_type is None:
        sect_type = parse_xml(f'<w:type {nsdecls("w")} w:val="continuous"/>')
        sectPr.insert(0, sect_type)
    else:
        sect_type.set(qn("w:val"), "continuous")

    # -- Set up bullet numbering in the document --
    # We need to add an abstractNum and num to the numbering part
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    # Create abstract numbering definition for bullets
    abstract_num_xml = (
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="10">'
        f'  <w:nsid w:val="AABB0001"/>'
        f'  <w:multiLevelType w:val="hybridMultilevel"/>'
        f'  <w:tmpl w:val="AABB0002"/>'
        f'  <w:lvl w:ilvl="0">'
        f'    <w:start w:val="1"/>'
        f'    <w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="\u00B7"/>'
        f'    <w:lvlJc w:val="left"/>'
        f'    <w:pPr>'
        f'      <w:ind w:left="720" w:hanging="360"/>'
        f'    </w:pPr>'
        f'    <w:rPr>'
        f'      <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>'
        f'    </w:rPr>'
        f'  </w:lvl>'
        f'</w:abstractNum>'
    )
    abstract_num = parse_xml(abstract_num_xml)
    # Insert abstractNum before any num elements
    first_num = numbering_elm.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elm.append(abstract_num)

    # Create num element referencing our abstract
    num_xml = (
        f'<w:num {nsdecls("w")} w:numId="1">'
        f'  <w:abstractNumId w:val="10"/>'
        f'</w:num>'
    )
    num_elm = parse_xml(num_xml)
    numbering_elm.append(num_elm)

    # =========================================================================
    # DOCUMENT CONTENT
    # =========================================================================

    # Fix zoom in settings (validation requires w:percent attribute)
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    # --- NAME ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_name, before=0, after=0, line=360, line_rule="auto")
    r = p_name.add_run("EZRA J MILLER")
    set_font(r, bold=True)

    # --- Contact Line 1 ---
    p_contact1 = doc.add_paragraph()
    p_contact1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_contact1, before=0, after=0)
    r = p_contact1.add_run("US & UK Dual Citizen · linkedin.com/in/ezrajmiller")
    set_font(r, bold=False)

    # --- Contact Line 2 ---
    p_contact2 = doc.add_paragraph()
    p_contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_contact2, before=0, after=0)
    r = p_contact2.add_run("+1 (207) 504-6622 · ezmiller@syr.edu · ezrajmiller@hotmail.com")
    set_font(r, bold=False)

    # --- Empty line ---
    add_empty_paragraph(doc)

    # =========================================================================
    # EDUCATION
    # =========================================================================
    add_section_header(doc, "EDUCATION")

    # Syracuse University line with right-tab date
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_right_tab(p, 10710)
    r1 = p.add_run("Syracuse University,")
    set_font(r1, bold=True)
    r2 = p.add_run(" School of Information Studies · College of Arts and Sciences")
    set_font(r2, bold=False)
    add_tab_run(p)
    r3 = p.add_run("May 2026")
    set_font(r3, bold=False)

    # Degrees line
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r = p.add_run("B.S. Applied Data Analytics · B.A. Psychology · Minor in Business Management")
    set_font(r, bold=False)

    # Relevant Coursework
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r1 = p.add_run("Relevant Coursework: ")
    set_font(r1, bold=True)
    r2 = p.add_run(
        "Big Data Analytics | Data-Driven Inquiry | Advanced Big Data Management "
        "| Building Human-Centered AI Applications | Information Visualization"
    )
    set_font(r2, bold=False)

    # Empty line
    add_empty_paragraph(doc)

    # =========================================================================
    # PROFESSIONAL EXPERIENCE
    # =========================================================================
    add_section_header(doc, "PROFESSIONAL EXPERIENCE")

    # -- Co-Owner, Comstock Customs --
    add_entry_header(doc, "Co-Owner, ", "Comstock Customs, Syracuse NY", "April 2025 \u2013 Present")

    add_bullet(doc, "Design all artwork using Adobe Photoshop and Illustrator")
    add_bullet(doc, "Coordinate with clients and negotiate specifications")
    add_bullet(doc, "Build automated systems for customer contact and order management using Excel")
    add_bullet(doc, "Optimize physical screen-printing process to produce 120 shirts per hour, reducing manual labor")
    add_bullet(doc, "Manage finances, inventory, and operational logistics")
    add_bullet(doc, "Achieved 77% gross profit margin and $3,000+ in revenue")

    # Empty line
    add_empty_paragraph(doc)

    # -- Webmaster, Sigma Chi --
    add_entry_header(doc, "Webmaster, ", "Sigma Chi Fraternity, Syracuse NY", "August 2024 \u2013 May 2025")

    add_bullet(doc, "Directed recruitment marketing across Instagram and the chapter website, increasing chapter visibility")
    add_bullet(doc, "Led philanthropy promotion through digital engagement and merchandise sales, contributing to chapter\u2019s record-breaking $155,000 fundraising effort")

    # Empty line
    add_empty_paragraph(doc)

    # =========================================================================
    # PROJECTS
    # =========================================================================
    add_section_header(doc, "PROJECTS")

    # -- Music Library Digitization --
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_right_tab(p, 10710)
    r1 = p.add_run("Music Library Digitization")
    set_font(r1, bold=True)
    add_tab_run(p)
    r2 = p.add_run("January 2026")
    set_font(r2, bold=False)

    add_bullet(doc, "Built ETL pipeline digitizing a CD library into 1.5TB archive of over 50,000 lossless audio tracks")
    add_bullet(doc, "Automated conversion and tagging workflow using bash scripting to minimize manual processing time")
    add_bullet(doc, "Utilized web scraping system to ensure metadata veracity")
    add_bullet(doc, "Optimized storage through WAV-to-FLAC compression while maintaining fidelity")

    # Empty line
    add_empty_paragraph(doc)

    # -- City of Syracuse Service Requests Analysis --
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_right_tab(p, 10710)
    r1 = p.add_run("City of Syracuse Service Requests Analysis")
    set_font(r1, bold=True)
    add_tab_run(p)
    r2 = p.add_run("April 2025")
    set_font(r2, bold=False)

    add_bullet(doc, "Collaborated with a team to analyze 60,000+ municipal service requests using R")
    add_bullet(doc, "Identified drivers behind response times and resolution efficiency")
    add_bullet(doc, "Engineered features including response-time metrics and mobile reporting indicators to support association rules mining and linear regression modeling")
    add_bullet(doc, "Built geospatial visualizations mapping request density and closure speed across Syracuse zip codes")

    # Empty line
    add_empty_paragraph(doc)

    # =========================================================================
    # LEADERSHIP EXPERIENCE
    # =========================================================================
    add_section_header(doc, "LEADERSHIP EXPERIENCE")

    # -- Lead Guitarist, Oldenvice --
    add_entry_header(doc, "Lead Guitarist, ", "Oldenvice, Rome NY", "February 2023 \u2013 Present")

    add_bullet(doc, "Lead 4-person band with regular performances at Central NY venues; developed 2-hour repertoire")
    add_bullet(doc, "Direct rehearsals, coordinate performance execution, and manage ensemble cohesion")
    add_bullet(doc, "Recorded and released multiple singles; built dedicated local following")

    # Empty line
    add_empty_paragraph(doc)

    # =========================================================================
    # TECHNICAL SKILLS
    # =========================================================================
    add_section_header(doc, "TECHNICAL SKILLS")

    # Languages
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r1 = p.add_run("Languages: ")
    set_font(r1, bold=True)
    r2 = p.add_run("Python, SQL, R, HTML/CSS/JS")
    set_font(r2, bold=False)

    # Tools & Software
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r1 = p.add_run("Tools & Software:")
    set_font(r1, bold=True)
    r2 = p.add_run(" MS Word/Excel, GitHub, Docker, Jupyter, VSCode, Adobe Photoshop/Illustrator, Logic Pro X, Linux, Bash")
    set_font(r2, bold=False)

    # Database Management
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r1 = p.add_run("Database Management: ")
    set_font(r1, bold=True)
    r2 = p.add_run("Apache Spark, Hadoop, Hive, Drill, Redis, MinIO, Neo4j")
    set_font(r2, bold=False)

    # =========================================================================
    # SAVE
    # =========================================================================
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Ezra_Miller_Resume_Data.docx")
    doc.save(output_path)
    print(f"Resume saved to {output_path}")


if __name__ == "__main__":
    build_resume()