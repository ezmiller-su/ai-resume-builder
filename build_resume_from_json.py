#!/usr/bin/env python3
"""
Builds a resume .docx from a JSON data file using python-docx.

Usage:
    python build_resume.py                      # uses resume_data.json in same dir
    python build_resume.py my_data.json         # uses a custom JSON file
    python build_resume.py data.json out.docx   # custom input and output
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json
import os
import sys


# =============================================================================
# Low-level formatting helpers
# =============================================================================

def set_font(run, name="Times New Roman", size=Pt(12), bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _get_or_create_pPr(paragraph):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        paragraph._element.insert(0, pPr)
    return pPr


_PPR_ORDER = [
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
    "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
    "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens",
    "w:kinsoku", "w:wordWrap", "w:overflowPunct", "w:topLinePunct",
    "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing",
    "w:mirrorIndents", "w:suppressOverlap", "w:jc", "w:textDirection",
    "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl", "w:divId",
    "w:cnfStyle", "w:rPr",
]


def _insert_pPr_child(pPr, child_elem):
    child_tag = child_elem.tag
    child_local = child_tag.split("}")[-1] if "}" in child_tag else child_tag
    child_qname = f"w:{child_local}"
    try:
        child_idx = _PPR_ORDER.index(child_qname)
    except ValueError:
        pPr.append(child_elem)
        return
    for existing in pPr:
        existing_tag = existing.tag
        existing_local = existing_tag.split("}")[-1] if "}" in existing_tag else existing_tag
        existing_qname = f"w:{existing_local}"
        try:
            existing_idx = _PPR_ORDER.index(existing_qname)
        except ValueError:
            continue
        if existing_idx > child_idx:
            existing.addprevious(child_elem)
            return
    pPr.append(child_elem)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_rule=None):
    pPr = _get_or_create_pPr(paragraph)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = parse_xml(f"<w:spacing {nsdecls('w')}/>")
        _insert_pPr_child(pPr, spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line is not None:
        spacing.set(qn("w:line"), str(line))
    if line_rule is not None:
        spacing.set(qn("w:lineRule"), line_rule)


def add_bottom_border(paragraph):
    pPr = _get_or_create_pPr(paragraph)
    _insert_pPr_child(pPr, parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'))


def add_right_tab(paragraph, pos=10710):
    pPr = _get_or_create_pPr(paragraph)
    _insert_pPr_child(pPr, parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        f'  <w:tab w:val="right" w:pos="{pos}"/>'
        f'</w:tabs>'))


def add_tab_run(paragraph):
    r = parse_xml(f'<w:r {nsdecls("w")}><w:tab/></w:r>')
    paragraph._element.append(r)
    return r


# =============================================================================
# Paragraph-level emitters
# =============================================================================

def emit_bullet(doc, text):
    """Emit a single bullet-point paragraph."""
    p = doc.add_paragraph()
    pPr = _get_or_create_pPr(p)
    _insert_pPr_child(pPr, parse_xml(
        f'<w:pStyle {nsdecls("w")} w:val="ListParagraph"/>'))
    _insert_pPr_child(pPr, parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="0"/><w:numId w:val="1"/>'
        f'</w:numPr>'))
    set_paragraph_spacing(p, before=0, after=0)
    r = p.add_run(text)
    set_font(r, bold=False)


def emit_empty(doc):
    """Emit a blank spacer paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r = p.add_run()
    set_font(r, bold=False)


def emit_section_header(doc, text):
    """Emit a bold section title with a bottom rule (e.g. EDUCATION)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_bottom_border(p)
    r = p.add_run(text)
    set_font(r, bold=True)


def emit_dated_line(doc, bold_text, normal_text, date_text):
    """Emit: **bold_text** normal_text  →tab→  date_text."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    add_right_tab(p)
    r1 = p.add_run(bold_text)
    set_font(r1, bold=True)
    if normal_text:
        r2 = p.add_run(normal_text)
        set_font(r2, bold=False)
    tab_r = add_tab_run(p)
    tab_r.insert(0, parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'  <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'  <w:color w:val="000000"/>'
        f'</w:rPr>'))
    r3 = p.add_run(date_text)
    set_font(r3, bold=False)


def emit_label_value(doc, label, value):
    """Emit: **label** value  (one plain line)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    r1 = p.add_run(label)
    set_font(r1, bold=True)
    r2 = p.add_run(value)
    set_font(r2, bold=False)


# =============================================================================
# Section renderers — one per JSON section "type"
# =============================================================================

def render_education(doc, entries):
    for entry in entries:
        emit_dated_line(doc,
                        f'{entry["institution_bold"]} ',
                        entry.get("institution_detail", ""),
                        entry["date"])
        if "degree" in entry:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=0)
            r = p.add_run(entry["degree"])
            set_font(r, bold=False)
        if "coursework" in entry:
            cw = entry["coursework"]
            cw_text = " \u00b7 ".join(cw) if isinstance(cw, list) else cw
            emit_label_value(doc,
                             entry.get("coursework_label", "Relevant Coursework: "),
                             cw_text)


def render_experience(doc, entries):
    for i, entry in enumerate(entries):
        if i > 0:
            emit_empty(doc)
        emit_dated_line(doc, f'{entry["role"]} ', entry["org"], entry["date"])
        for bullet in entry.get("bullets", []):
            emit_bullet(doc, bullet)


def render_projects(doc, entries):
    for i, entry in enumerate(entries):
        if i > 0:
            emit_empty(doc)
        emit_dated_line(doc, f'{entry["name"]} ', None, entry["date"])
        for bullet in entry.get("bullets", []):
            emit_bullet(doc, bullet)


def render_skills(doc, entries):
    for entry in entries:
        value = entry.get("value") or " \u00b7 ".join(entry.get("items", []))
        emit_label_value(doc, entry["label"], value)


_RENDERERS = {
    "education":  render_education,
    "experience": render_experience,
    "projects":   render_projects,
    "skills":     render_skills,
}


# =============================================================================
# Document scaffolding (styles, page setup, numbering)
# =============================================================================

def init_document():
    """Create a blank doc with correct styles, page size, and bullet numbering."""
    doc = Document()

    # --- Normal style ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")

    # --- List Paragraph style ---
    lp = doc.styles["List Paragraph"]
    lp.font.name = "Times New Roman"
    lp.font.size = Pt(12)
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)

    # --- Page setup: US Letter, 0.5" margins ---
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sectPr = sec._sectPr
    st = sectPr.find(qn("w:type"))
    if st is None:
        sectPr.insert(0, parse_xml(
            f'<w:type {nsdecls("w")} w:val="continuous"/>'))
    else:
        st.set(qn("w:val"), "continuous")

    # --- Bullet numbering ---
    numbering_elm = doc.part.numbering_part.numbering_definitions._numbering
    abstract = parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="10">'
        f'  <w:nsid w:val="AABB0001"/>'
        f'  <w:multiLevelType w:val="hybridMultilevel"/>'
        f'  <w:tmpl w:val="AABB0002"/>'
        f'  <w:lvl w:ilvl="0">'
        f'    <w:start w:val="1"/>'
        f'    <w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="\u00B7"/>'
        f'    <w:lvlJc w:val="left"/>'
        f'    <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
        f'    <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol"'
        f'           w:hint="default"/></w:rPr>'
        f'  </w:lvl>'
        f'</w:abstractNum>')
    first_num = numbering_elm.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        numbering_elm.append(abstract)
    numbering_elm.append(parse_xml(
        f'<w:num {nsdecls("w")} w:numId="1">'
        f'  <w:abstractNumId w:val="10"/>'
        f'</w:num>'))

    # --- Fix zoom validation ---
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    return doc


# =============================================================================
# Contact-line balancing
# =============================================================================

_SEP = " \u00b7 "
# Approximate max characters that fit on one centered line
# (Times New Roman 12pt, 7.5" content width ≈ 80 chars)
_MAX_LINE_CHARS = 80


def _balance_contact_lines(items):
    """Return 1 or 2 joined strings, splitting only if a single line overflows."""
    if not items:
        return []
    # Flatten if wrapped in an outer list: [["a","b","c"]] → ["a","b","c"]
    if items and isinstance(items[0], list):
        items = [x for sub in items for x in sub]
    single = _SEP.join(items)
    if len(single) <= _MAX_LINE_CHARS:
        return [single]
    # Split as evenly as possible across two lines
    best_split, best_diff = 1, float("inf")
    for i in range(1, len(items)):
        a = _SEP.join(items[:i])
        b = _SEP.join(items[i:])
        diff = abs(len(a) - len(b))
        if diff < best_diff:
            best_diff = diff
            best_split = i
    return [_SEP.join(items[:best_split]), _SEP.join(items[best_split:])]


# =============================================================================
# Main: load JSON → render → save
# =============================================================================

def build_resume(data, output_path):
    doc = init_document()

    # ---- Header ----
    header = data["header"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line=360, line_rule="auto")
    r = p.add_run(header["name"])
    set_font(r, bold=True)

    for line_text in _balance_contact_lines(header.get("contact_lines", [])):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0)
        r = p.add_run(line_text)
        set_font(r, bold=False)

    emit_empty(doc)

    # ---- Professional summary (optional) ----
    summary = header.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0)
        r = p.add_run(summary)
        set_font(r, bold=False)
        emit_empty(doc)

    # ---- Sections ----
    for i, section in enumerate(data["sections"]):
        emit_section_header(doc, section["title"])

        renderer = _RENDERERS.get(section["type"])
        if renderer is None:
            raise ValueError(f"Unknown section type: {section['type']!r}")
        renderer(doc, section["entries"])

        # Blank line between sections (not after the last one)
        if i < len(data["sections"]) - 1:
            emit_empty(doc)

    doc.save(output_path)
    print(f"Resume saved to {output_path}")


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))

    json_path = (sys.argv[1] if len(sys.argv) > 1
                 else os.path.join(script_dir, "ejm_info.json"))
    docx_path = (sys.argv[2] if len(sys.argv) > 2
                 else os.path.join(script_dir, "Ezra_Miller_Resume_Data.docx"))

    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    build_resume(data, docx_path)